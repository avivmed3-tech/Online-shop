# -*- coding: utf-8 -*-
"""יצירת קובץ Word מעוצב של הצעת מחיר לאפליקציית הקטלוג B2B (RTL מלא)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ===== פלטת צבעים =====
NAVY = RGBColor(0x16, 0x2A, 0x4A)      # כהה לכותרות
ACCENT = RGBColor(0x2E, 0x86, 0xC1)    # תכלת מודגש
GREY = RGBColor(0x5A, 0x5A, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = "EAF2F8"                        # רקע בהיר לטבלה
ZEBRA = "F4F8FB"
NAVY_HEX = "162A4A"
ACCENT_HEX = "2E86C1"

FONT = "Arial"

# ===== סדר הסכמה של OOXML (חובה כדי שהקובץ ייטען בכל עורך) =====
PPR_ORDER = ["w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
             "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
             "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
             "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
             "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
             "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
             "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
             "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
             "w:rPr", "w:sectPr", "w:pPrChange"]
TBLPR_ORDER = ["w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
               "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW",
               "w:jc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders",
               "w:tblShd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
               "w:tblCaption", "w:tblDescription", "w:tblPrChange"]
RPR_ORDER = ["w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
             "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
             "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
             "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern",
             "w:position", "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect",
             "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl", "w:cs",
             "w:em", "w:lang", "w:eastAsianLayout", "w:specVanish", "w:oMath"]
TCPR_ORDER = ["w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge",
              "w:tcBorders", "w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
              "w:tcFitText", "w:vAlign", "w:hideMark", "w:cellIns", "w:cellDel",
              "w:cellMerge", "w:tcPrChange"]
SECTPR_ORDER = ["w:headerReference", "w:footerReference", "w:footnotePr",
                "w:endnotePr", "w:type", "w:pgSz", "w:pgMar", "w:paperSrc",
                "w:pgBorders", "w:lnNumType", "w:pgNumType", "w:cols",
                "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
                "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
                "w:printerSettings", "w:sectPrChange"]


def insert_in_order(parent, child, order):
    """מכניס אלמנט למקום הנכון לפי סדר הסכמה."""
    tag = child.tag.split("}")[-1]
    idx = order.index("w:" + tag)
    after = order[idx + 1:]
    for existing in parent:
        et = existing.tag.split("}")[-1]
        if ("w:" + et) in after:
            existing.addprevious(child)
            return
    parent.append(child)


doc = Document()

# ===== שוליים =====
sec = doc.sections[0]
sec.top_margin = Cm(1.6)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

# ===== ברירת מחדל: RTL בכל המסמך =====
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:cs"), FONT)
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)


def _section_rtl():
    """הופך את כיוון הסעיף ל-RTL."""
    bidi = OxmlElement("w:bidi")
    insert_in_order(sec._sectPr, bidi, SECTPR_ORDER)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    insert_in_order(pPr, bidi, PPR_ORDER)


def run_rtl(run):
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    insert_in_order(rPr, rtl, RPR_ORDER)


def add_par(text="", size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT,
            space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        run_rtl(r)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = FONT
        if color is not None:
            r.font.color.rgb = color
    return p


def para_shade(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    insert_in_order(pPr, shd, PPR_ORDER)


def para_border(paragraph, edge="bottom", color="162A4A", sz="12", space="1"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        insert_in_order(pPr, pbdr, PPR_ORDER)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pbdr.append(el)


def table_rtl(table):
    """כיוון טבלה מימין לשמאל — העמודה הראשונה מופיעה מימין."""
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    insert_in_order(tblPr, bidi, TBLPR_ORDER)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    insert_in_order(tcPr, shd, TCPR_ORDER)


def cell_valign(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    insert_in_order(tcPr, va, TCPR_ORDER)


def cell_text(cell, text, bold=False, color=None, size=11,
              align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    run_rtl(r)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = FONT
    if color is not None:
        r.font.color.rgb = color
    cell_valign(cell)


_section_rtl()

# ============================================================
#  באנר כותרת (טבלה ברוחב מלא עם רקע כהה)
# ============================================================
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner.autofit = False
bcell = banner.rows[0].cells[0]
shade_cell(bcell, NAVY_HEX)
bcell.width = Cm(17.4)
bcell.text = ""

bp = bcell.paragraphs[0]
set_rtl(bp)
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(14)
bp.paragraph_format.space_after = Pt(2)
r = bp.add_run("הצעת מחיר")
run_rtl(r)
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = WHITE
r.font.name = FONT

bp2 = bcell.add_paragraph()
set_rtl(bp2)
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp2.paragraph_format.space_before = Pt(0)
bp2.paragraph_format.space_after = Pt(14)
r2 = bp2.add_run("אפליקציית קטלוג והזמנות עסקי  |  B2B")
run_rtl(r2)
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0xBF, 0xD7, 0xEA)
r2.font.name = FONT

add_par("", space_after=8)

# ===== פרטי ההצעה =====
info = doc.add_table(rows=2, cols=2)
table_rtl(info)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_text(info.rows[0].cells[0], "תאריך:", bold=True, color=NAVY, size=11)
cell_text(info.rows[0].cells[1], date.today().strftime("%d/%m/%Y"), color=GREY)
cell_text(info.rows[1].cells[0], "לכבוד:", bold=True, color=NAVY, size=11)
cell_text(info.rows[1].cells[1], "___________________", color=GREY)
for row in info.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(15)

add_par("", space_after=6)

# ===== תיאור הפרויקט =====
h = add_par("תיאור הפרויקט", size=15, bold=True, color=NAVY, space_before=4, space_after=6)
para_border(h, "bottom", ACCENT_HEX, "8", "2")

add_par(
    "פיתוח והפעלה של אפליקציית קטלוג הזמנות עסקית (B2B), בה כל לקוח עסקי "
    "מתחבר עם משתמש אישי וצופה במחירים המותאמים לו (לפי אחוז הנחה), מבצע "
    "הזמנות הנשמרות במערכת, ובעל העסק מנהל מוצרים, לקוחות והזמנות מלוח בקרה ייעודי.",
    size=11, space_after=10, line=1.3)

add_par("המערכת כוללת:", size=12.5, bold=True, color=NAVY, space_after=5)
for item in [
    "אפליקציית ווב מותקנת (PWA) הנפתחת מהדפדפן ומהנייד.",
    "מערכת התחברות ומשתמשים עם הרשאות (בעל עסק / לקוח).",
    "ניהול מוצרים, לקוחות והגדרת אחוז הנחה אישי לכל לקוח.",
    "מחירים אישיים לכל לקוח וקבלת הזמנות ישירות לעסק.",
    "אבטחת מידע ותמחור בצד השרת (כל לקוח רואה רק את הנתונים שלו).",
    "אחסון וענן מנוהל (Supabase + GitHub Pages).",
]:
    p = add_par(size=11, space_after=4, line=1.2)
    rb = p.add_run("◄  ")
    run_rtl(rb)
    rb.font.size = Pt(9)
    rb.font.color.rgb = ACCENT
    rt = p.add_run(item)
    run_rtl(rt)
    rt.font.size = Pt(11)
    rt.font.name = FONT
    p.paragraph_format.right_indent = Cm(0.4)

# ============================================================
#  טבלת מחיר
# ============================================================
h2 = add_par("פירוט מחיר", size=15, bold=True, color=NAVY, space_before=14, space_after=8)
para_border(h2, "bottom", ACCENT_HEX, "8", "2")

table = doc.add_table(rows=1, cols=3)
table_rtl(table)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
table.autofit = False

widths = [Cm(5.0), Cm(7.9), Cm(4.5)]

# כותרת
hdr = table.rows[0].cells
for c, txt, w in zip(hdr, ["שירות", "תיאור", "מחיר (כולל מע\"מ)"], widths):
    cell_text(c, txt, bold=True, color=WHITE, size=11.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(c, NAVY_HEX)
    c.width = w

rows = [
    ("בנייה והרצת האפליקציה", "פיתוח, הקמה והעלאה לאוויר של אפליקציית הקטלוג", "3,000 ש\"ח"),
    ("ניהול חודשי", "תחזוקה, אחסון, גיבוי ותמיכה שוטפת", "250 ש\"ח / חודש"),
]
for i, (service, desc, price) in enumerate(rows):
    cells = table.add_row().cells
    cell_text(cells[0], service, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.RIGHT)
    cell_text(cells[1], desc, align=WD_ALIGN_PARAGRAPH.RIGHT, color=GREY)
    cell_text(cells[2], price, bold=True, color=ACCENT, size=11.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for c, w in zip(cells, widths):
        c.width = w
    if i % 2 == 1:
        for c in cells:
            shade_cell(c, ZEBRA)

# שורת סיכום — תשלום חד-פעמי להקמה
sum_cells = table.add_row().cells
merged = sum_cells[0].merge(sum_cells[1])
cell_text(merged, "תשלום חד-פעמי להקמה", bold=True, color=NAVY, size=11.5,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(sum_cells[2], "3,000 ש\"ח", bold=True, color=NAVY, size=12,
          align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(merged, LIGHT)
shade_cell(sum_cells[2], LIGHT)

# שורת התחייבות חודשית
mon_cells = table.add_row().cells
merged2 = mon_cells[0].merge(mon_cells[1])
cell_text(merged2, "עלות חודשית שוטפת", bold=True, color=NAVY, size=11.5,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(mon_cells[2], "250 ש\"ח", bold=True, color=NAVY, size=12,
          align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(merged2, LIGHT)
shade_cell(mon_cells[2], LIGHT)

# ===== הערות =====
h3 = add_par("הערות", size=15, bold=True, color=NAVY, space_before=16, space_after=6)
para_border(h3, "bottom", ACCENT_HEX, "8", "2")
for note in [
    "המחירים נקובים בשקלים חדשים (ש\"ח).",
    "דמי הניהול החודשיים נגבים החל מהחודש שלאחר עליית האפליקציה לאוויר.",
    "ההצעה כוללת ליווי והדרכה ראשונית לשימוש במערכת.",
    "תוקף ההצעה: 30 יום מתאריך הוצאתה.",
]:
    p = add_par(size=10.5, space_after=4, line=1.2)
    rb = p.add_run("•  ")
    run_rtl(rb)
    rb.font.color.rgb = ACCENT
    rt = p.add_run(note)
    run_rtl(rt)
    rt.font.size = Pt(10.5)
    rt.font.color.rgb = GREY
    rt.font.name = FONT
    p.paragraph_format.right_indent = Cm(0.4)

# ===== חתימה =====
add_par("בברכה,", size=11, bold=True, color=NAVY, space_before=20, space_after=2)
add_par("___________________", size=11, color=GREY, space_after=2)
add_par("חתימה ותאריך", size=10, color=GREY)

doc.save("/home/user/Online-shop/הצעת_מחיר.docx")
print("saved")
